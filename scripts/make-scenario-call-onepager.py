#!/usr/bin/env python3
"""Generate the one-page Word note on the IAM scenario submission call.

Output: project-documents/2026-06-12 Scenario Submission Call - One-pager.docx
Requires the figure produced by scripts/make-scenario-call-figure.py.
Re-run after editing the content below: python3 scripts/make-scenario-call-onepager.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x2F, 0x6E, 0x5B)  # MethodHub sector green
GREY = RGBColor(0x55, 0x5B, 0x63)

doc = Document()

# One page: narrow margins, compact base style.
for section in doc.sections:
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(2)


def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    return p


def para(text, italic=False, size=9, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(lead, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(lead + " — " if lead else "")
    r.bold = True
    p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(9)
    return p


# ── Title ────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("An open scenario submission call for IAM modellers")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ACCENT
t.paragraph_format.space_after = Pt(1)

s = doc.add_paragraph()
r = s.add_run(
    "Our thinking on scenarios: what we need them for, how to acquire them, "
    "the workflow on our side, and the risks — ESABCC MethodHub, draft for discussion, 12 June 2026"
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY
s.paragraph_format.space_after = Pt(6)

# ── What we need scenarios for ───────────────────────────────────────────────
heading("What do we need scenarios for? Two sides of the same coin")
bullet("Checking the benchmarks, adding context (Fig. 1a–b)",
       "Our monitoring frameworks tell us what to track but not how fast it must move. The advisory board's sector "
       "flow charts (Figures 12–64 of 'Towards EU climate neutrality') carry ~40 progress indicators; today these "
       "are read against observed trends, Member State WEM/WAM projections and EC benchmarks only. Matched scenario "
       "corridors turn every indicator into a quantified benchmark — is solar deployment, renovation, or the cattle "
       "herd inside the corridor that EU-climate-neutrality pathways require? — allow us to assess whether EC "
       "scenarios are too optimistic or too pessimistic against the broader scientific landscape, and keep our gap "
       "assessments evidence-based as policy and model vintages evolve.")
bullet("Understanding the inherent dynamics of the transition (Fig. 1c)",
       "Meeting the 2030 target is not the same as being on track for 2050. Sectors bank the easy gains first — in "
       "energy, the low-hanging fruit of mature renewables roll-out. What remains is the hard core: structural and "
       "technological constraints (limited availability of CCS, hydrogen, critical materials, infrastructure) and "
       "major societal challenges (diets, consumption, demand). A call that includes sectoral modelling outputs "
       "builds a bottom-up view of these system dynamics, so we can see where we really stand beyond the headline "
       "GHG indicator — reading between the lines of apparent target compliance.")

# ── Options ──────────────────────────────────────────────────────────────────
heading("Which options do we have to acquire them?")
bullet("Reuse existing databases", "AR6 / ECEMF / Scenario Explorer snapshots. Cheap, but vintages age, EU27 "
       "sectoral detail is thin, and we cannot ask for missing variables.")
bullet("Commission modelling", "Full control over scenario design, but slow, costly, and yields few models — weak "
       "ensembles and a perception of hand-picked evidence.")
bullet("Open submission call (proposed)", "Invite IAM and sectoral modelling teams to submit economy-wide or "
       "sectoral EU scenarios directly to the MethodHub. Best ensemble breadth and freshness at the lowest cost; "
       "the other two options remain available as fallback and complement.")

# ── Figure ───────────────────────────────────────────────────────────────────
fig_p = doc.add_paragraph()
fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_p.paragraph_format.space_before = Pt(3)
fig_p.paragraph_format.space_after = Pt(0)
fig_p.add_run().add_picture("project-documents/assets/scenario-call-logic-figure.png", width=Cm(15.5))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap.add_run(
    "Figure 1 | The two purposes of the call (schematic). a, Current assessment logic. b, The submitted ensemble "
    "(band: range; line: median) makes official projections and EC benchmarks testable. c, The gap to the net-zero "
    "pathway is closed first by mature-technology roll-out, later by structural and societal change. d1–d3, zooming "
    "into that hard core with one sectoral-model example per sector: the industry electrolyser capacity gap, "
    "transport biomass demand vs sustainable availability, and buildings heat-pump roll-out slowed by the old, "
    "unrenovated stock."
)
r.italic = True
r.font.size = Pt(7.5)
r.font.color.rgb = GREY
cap.paragraph_format.space_after = Pt(2)

# ── How to read the figure ───────────────────────────────────────────────────
heading("What the figure shows")
bullet("Panel a (today)", "Both forward-looking elements come from official sources only: the Member State "
       "WEM/WAM projections (reported under the Governance Regulation, aggregated by the EEA) answer 'will we meet "
       "the benchmark under current and planned policies?', and the EC-scenario benchmarks set the goal posts. "
       "There is no independent scientific check and no uncertainty range around either.")
bullet("Panel b (checking & context)", "Both elements are kept, but the submitted scenario ensemble is added "
       "alongside: the shaded band is the range across all submitted runs, the solid line their median. Where the "
       "official projections and benchmarks sit relative to the band shows at a glance whether they are optimistic "
       "or conservative against the scientific landscape.")
bullet("Panel c (transition dynamics)", "Hitting a near-term target says little about the road beyond it. The gap "
       "between the current-effort baseline and the net-zero pathway is small in 2030 and closed mostly by "
       "mature-technology roll-out — the low-hanging fruit, which delivers early and then saturates. The much "
       "larger 2050 gap must come from the hard core: structural and societal change (technology availability "
       "limits, infrastructure, diets, consumption, demand). Sectoral-model submissions resolve this split per "
       "sector, bottom-up, beyond the GHG headline.")
bullet("Panels d1–d3 (the sectoral zoom)", "One example per sector of the detailed outputs we want submitted: "
       "industry — the gap between the current electrolyser deployment trend and what pathways require; transport — "
       "biofuel demand running into the sustainable biomass supply that land-use models can certify, with "
       "competition across sectors; buildings — the heat-pump roll-out required versus a trend slowed by the old, "
       "unrenovated stock. This is the 'reading between the lines': each panel turns an invisible structural "
       "constraint into a measurable gap.")

# ── Workflow ─────────────────────────────────────────────────────────────────
heading("What would the workflow be on our side?")
bullet("Submission infrastructure", "Built on the MethodHub, which already bundles a Scenario Explorer snapshot "
       "(63 EU27 runs from 9 model versions) and the indicator database. Teams submit IAMC-template files through "
       "a portal; no new platform is needed.")
bullet("Automated matching", "Each submitted run is resolved against the indicator framework — a prototype flow "
       "chart ('Scenario call — IAM-matched indicators') already maps every indicator of the default board to an "
       "IAMC variable and submission track (economy-wide / sectoral / new variable requested).")
bullet("Filtering & analysis", "A fully-fledged vetting, filtering and corridor-analysis toolset is built into the "
       "MethodHub pipeline (harmonisation diagnostics as used for the existing snapshot), so screening submissions "
       "is largely automatic.")
bullet("Distribution", "The main genuine effort. With two leading IAM modellers on the board — Keywan Riahi "
       "(MESSAGEix-GLOBIOM, IIASA) and Detlef van Vuuren (IMAGE, PBL) — the call travels through the IAMC, ECEMF "
       "and NAVIGATE/ENGAGE networks at essentially no cost and with high credibility.")

# ── Benefits ─────────────────────────────────────────────────────────────────
heading("Benefits")
bullet("Quantified monitoring", "Every indicator gains a scenario corridor; assessments shift from directional to "
       "quantitative, continuously refreshable.")
bullet("Transition dynamics, bottom-up", "Sectoral submissions show how much of the remaining effort still rests "
       "on mature-technology roll-out versus hard structural and societal change — so early target-hitting cannot "
       "hide a hard residual.")
bullet("Broader, fresher evidence", "An open ensemble across many teams beats any single commissioned study; "
       "sectoral submissions (fleet, building-stock, agro-economic models) fill known IAM blind spots such as "
       "renovation rates, mode shares and livestock intensities.")
bullet("Low marginal cost", "Tooling is an extension of existing MethodHub infrastructure; modellers' incentive is "
       "visibility of their scenarios in advisory-board assessments.")
bullet("Transparency", "A documented, open call with published criteria strengthens the board's independence "
       "compared with selecting scenarios ad hoc.")

# ── Risks ────────────────────────────────────────────────────────────────────
heading("Risks")
bullet("Low or skewed response", "The central risk. Mitigated by the board's network distribution, a low submission "
       "burden (standard IAMC template), and seeding the database with the 63 runs we already hold so the tool is "
       "useful from day one.")
bullet("Quality and comparability", "Mitigated by automated vetting and harmonisation diagnostics, plus published "
       "inclusion criteria; submission does not imply endorsement.")
bullet("Coverage gaps", "Some indicators (circular material use, food waste, first-generation biofuel shares) are "
       "not standard outputs; the call flags them as requested variables and the sectoral track recruits the models "
       "that can report them.")
bullet("Process cost", "Concentrated in distribution and follow-up rather than in engineering — a bounded, "
       "plannable effort for the secretariat.")

# ── Bottom line ──────────────────────────────────────────────────────────────
heading("Bottom line")
para(
    "The call is feasible with modest effort: the infrastructure and analysis tooling can be delivered on the "
    "MethodHub, the indicator matching already exists in prototype, and the one real success factor — that teams "
    "actually submit — is exactly where the board is strongest, through Keywan's and Detlef's networks. We propose "
    "preparing a pilot call for autumn 2026."
)

out = "project-documents/2026-06-12 Scenario Submission Call - One-pager.docx"
doc.save(out)
print("wrote", out)
